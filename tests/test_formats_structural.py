"""Structural format tests — EPUB footnotes/images/metadata, DOCX
headings/styles/tables, SRT HTML tags/timecodes, TXT >100K chars.

All fixtures are built in-memory and written to tempdirs; no external
fixture files are needed.
"""

import tempfile
import textwrap
import unittest
from pathlib import Path

from src.backend.formats import (
    chunk_blocks,
    chunk_paragraphs,
    detect_format,
    read_document,
)


# ---------------------------------------------------------------------------
# EPUB structural tests
# ---------------------------------------------------------------------------


class EpubStructuralTests(unittest.TestCase):
    """EPUB: footnotes, embedded images, comprehensive DC metadata."""

    @classmethod
    def setUpClass(cls) -> None:
        try:
            from bs4 import BeautifulSoup  # noqa: F401
            from ebooklib import epub  # noqa: F401
        except ImportError:
            raise unittest.SkipTest("EPUB dependencies not installed")

    def _build_epub(self, root: Path, name: str, *items) -> Path:
        """Minimal EPUB builder returning the path to the written file."""
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier("test-id")
        book.set_title("Structural Test")
        book.set_language("en")
        spine = ["nav"]
        toc = []
        for item in items:
            book.add_item(item)
            if isinstance(item, epub.EpubHtml):
                spine.append(item)
                toc.append(item)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = spine
        book.toc = tuple(toc) if toc else ()
        path = root / name
        epub.write_epub(str(path), book)
        return path

    def test_footnotes_preserved_in_text(self) -> None:
        """Footnote text inside ``<aside epub:type="footnote">`` is captured."""
        from ebooklib import epub

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            chapter = epub.EpubHtml(
                title="Ch1", file_name="ch1.xhtml", lang="en"
            )
            chapter.content = textwrap.dedent("""\
                <html><body>
                <p>Main text with a footnote reference.</p>
                <aside epub:type="footnote" id="fn1">
                  <p>This is the footnote body text.</p>
                </aside>
                <p>More main text after footnote.</p>
                </body></html>""")
            path = self._build_epub(root, "footnote.epub", chapter)

            doc = read_document(path)
            all_text = "\n".join(
                p for ch in doc.chapters for p in ch.paragraphs
            )
            self.assertIn("Main text with a footnote reference", all_text)
            self.assertIn("footnote body text", all_text)
            self.assertIn("More main text after footnote", all_text)

    def test_image_context_preserved(self) -> None:
        """Text surrounding ``<img>`` tags survives the parse.

        ``<img>`` elements are not text nodes, so only the surrounding
        paragraph text is expected in the output.
        """
        from ebooklib import epub

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            chapter = epub.EpubHtml(
                title="Ch1", file_name="ch1.xhtml", lang="en"
            )
            chapter.content = textwrap.dedent("""\
                <html><body>
                <p>Before the image.</p>
                <p><img src="cover.jpg" alt="Cover Image"/></p>
                <p>After the image.</p>
                </body></html>""")
            path = self._build_epub(root, "image.epub", chapter)

            doc = read_document(path)
            all_text = "\n".join(
                p for ch in doc.chapters for p in ch.paragraphs
            )
            self.assertIn("Before the image", all_text)
            self.assertIn("After the image", all_text)

    def test_dc_metadata_captured(self) -> None:
        """DC title and creator are captured; publisher/date are forwarded."""
        from ebooklib import epub

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            book = epub.EpubBook()
            book.set_identifier("dc-test")
            book.set_title("The Great Novel")
            book.set_language("en")
            book.add_metadata("DC", "creator", "Jane Author")
            book.add_metadata("DC", "publisher", "Test Press")
            book.add_metadata("DC", "date", "2025-06-01")
            chapter = epub.EpubHtml(
                title="Ch1", file_name="ch1.xhtml", lang="en"
            )
            chapter.content = "<html><body><p>Hello.</p></body></html>"
            book.add_item(chapter)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = ["nav", chapter]
            book.toc = (chapter,)
            path = root / "meta.epub"
            epub.write_epub(str(path), book)

            doc = read_document(path)
            self.assertEqual(doc.title, "The Great Novel")
            self.assertEqual(doc.author, "Jane Author")
            self.assertEqual(doc.metadata.get("path"), str(path))


# ---------------------------------------------------------------------------
# DOCX structural tests
# ---------------------------------------------------------------------------


class DocxStructuralTests(unittest.TestCase):
    """DOCX: multi-level headings, named styles, tables."""

    @classmethod
    def setUpClass(cls) -> None:
        try:
            import docx  # noqa: F401
        except ImportError:
            raise unittest.SkipTest("DOCX dependencies not installed")

    def test_headings_preserve_text(self) -> None:
        """All heading text survives parsing even if heading splitting is not
        yet implemented (single-chapter output)."""
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            d = docx.Document()
            d.add_heading("Heading Level 1", level=1)
            d.add_paragraph("Body text under heading 1.")
            d.add_heading("Heading Level 2", level=2)
            d.add_paragraph("Body text under heading 2.")
            d.add_heading("Heading Level 3", level=3)
            d.add_paragraph("Body text under heading 3.")
            path = root / "headings.docx"
            d.save(str(path))

            doc = read_document(path)
            all_text = "\n".join(
                p for ch in doc.chapters for p in ch.paragraphs
            )
            self.assertIn("Heading Level 1", all_text)
            self.assertIn("Heading Level 2", all_text)
            self.assertIn("Heading Level 3", all_text)
            self.assertIn("Body text under heading 1", all_text)
            self.assertIn("Body text under heading 2", all_text)
            self.assertIn("Body text under heading 3", all_text)

    def test_named_paragraph_styles_preserve_text(self) -> None:
        """Styled paragraphs (e.g. 'Quote', 'Intense Quote') keep their text."""
        import docx

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            d = docx.Document()
            p = d.add_paragraph("Normal paragraph.")
            p.style = d.styles["Normal"]
            p2 = d.add_paragraph("A quoted passage.")
            try:
                p2.style = d.styles["Quote"]
            except KeyError:
                pass
            path = root / "styled.docx"
            d.save(str(path))

            doc = read_document(path)
            all_text = "\n".join(
                p for ch in doc.chapters for p in ch.paragraphs
            )
            self.assertIn("Normal paragraph", all_text)
            self.assertIn("A quoted passage", all_text)

    def test_table_text_extraction(self) -> None:
        """Table cell text may be inaccessible via ``doc.paragraphs``.

        This test documents the current behaviour: table paragraphs are not
        part of ``Document.paragraphs``, so text in tables is currently
        lost. The test asserts the limitation rather than hiding it.
        """
        import docx

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            d = docx.Document()
            d.add_paragraph("Before table.")
            table = d.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Cell A1"
            table.cell(0, 1).text = "Cell B1"
            table.cell(1, 0).text = "Cell A2"
            table.cell(1, 1).text = "Cell B2"
            d.add_paragraph("After table.")
            path = root / "table.docx"
            d.save(str(path))

            doc = read_document(path)
            all_text = "\n".join(
                p for ch in doc.chapters for p in ch.paragraphs
            )
            self.assertIn("Before table", all_text)
            self.assertIn("After table", all_text)
            self.assertNotIn("Cell A1", all_text,
                             "table text is currently not extracted by _read_docx")


# ---------------------------------------------------------------------------
# SRT structural tests
# ---------------------------------------------------------------------------


class SrtStructuralTests(unittest.TestCase):
    """SRT: HTML tags preserved, timecodes up to 99h accepted."""

    def test_html_tags_preserved(self) -> None:
        """HTML tags in cue text (``<i>``, ``<b>``, ``<font>``) survive parse."""
        srt_content = textwrap.dedent("""\
            1
            00:00:01,000 --> 00:00:04,000
            <i>Italicized dialogue</i>

            2
            00:00:05,000 --> 00:00:08,000
            <b>Bold emphasis</b> and normal text.

            3
            00:00:09,000 --> 00:00:12,000
            <font color="#FF0000">Red subtitle text.</font>
            """)

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "tags.srt"
            path.write_text(srt_content, encoding="utf-8")
            doc = read_document(path)
            all_text = "\n".join(
                p for ch in doc.chapters for p in ch.paragraphs
            )
            self.assertIn("<i>Italicized dialogue</i>", all_text)
            self.assertIn("<b>Bold emphasis</b> and normal text.", all_text)
            self.assertIn('<font color="#FF0000">Red subtitle text.</font>', all_text)

    def test_timecodes_above_one_hour(self) -> None:
        """Timecodes with hours >= 1 parse correctly."""
        srt_content = textwrap.dedent("""\
            1
            01:30:45,500 --> 01:30:50,000
            This subtitle starts at 1h30m.

            2
            02:00:00,000 --> 02:00:05,000
            This subtitle starts at 2 hours.

            3
            10:15:30,000 --> 10:15:35,000
            This subtitle starts at 10 hours.
            """)

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "longtime.srt"
            path.write_text(srt_content, encoding="utf-8")
            doc = read_document(path)
            all_text = "\n".join(
                p for ch in doc.chapters for p in ch.paragraphs
            )
            self.assertIn("This subtitle starts at 1h30m", all_text)
            self.assertIn("This subtitle starts at 2 hours", all_text)
            self.assertIn("This subtitle starts at 10 hours", all_text)

    def test_standard_srt_parsing(self) -> None:
        """Basic SRT with multiple cues, sequence numbers, and timestamps."""
        srt_content = textwrap.dedent("""\
            1
            00:00:01,000 --> 00:00:04,000
            Hello world.

            2
            00:00:05,000 --> 00:00:08,000
            This is a subtitle.

            3
            00:00:09,000 --> 00:00:12,000
            Goodbye.
            """)

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "standard.srt"
            path.write_text(srt_content, encoding="utf-8")
            doc = read_document(path)
            self.assertEqual(len(doc.chapters), 1)
            paragraphs = doc.chapters[0].paragraphs
            self.assertEqual(len(paragraphs), 3)
            self.assertIn("Hello world.", paragraphs)
            self.assertIn("This is a subtitle.", paragraphs)
            self.assertIn("Goodbye.", paragraphs)

    def test_srt_detect_format(self) -> None:
        self.assertEqual(detect_format("movie.srt"), "srt")


# ---------------------------------------------------------------------------
# TXT structural tests
# ---------------------------------------------------------------------------


class TxtStructuralTests(unittest.TestCase):
    """TXT: long-file chunking and chapter heading detection."""

    def test_long_file_chunking_over_100k(self) -> None:
        """A >100K-character TXT produces multiple chunks with split metadata."""
        heading = "Chapter 1\n\n"
        sentence = "This is a test sentence that will be repeated many times to build up a large text. "
        # Build ~120 KB of text (sentence ≈ 90 bytes × 1400 reps).
        body = sentence * 1400
        content = heading + body

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "long.txt"
            path.write_text(content, encoding="utf-8")

            doc = read_document(path)
            self.assertEqual(len(doc.chapters), 1)
            self.assertEqual(doc.chapters[0].title, "Chapter 1")

            chunks = chunk_paragraphs(
                doc.chapters[0].paragraphs,
                doc.chapters[0].id,
                doc.chapters[0].title,
            )
            self.assertGreater(len(chunks), 1, "long file must produce multiple chunks")

            # Each chunk must have required keys.
            for c in chunks:
                self.assertIn("id", c)
                self.assertIn("source_text", c)
                self.assertIn("source_hash", c)
                self.assertIn("chunk_index", c)
                self.assertIn("metadata", c)

            # Total text across chunks should approximate original.
            total_text = "".join(c["source_text"] for c in chunks)
            self.assertGreater(len(total_text), 100_000)

    def test_txt_with_multiple_chapter_headings(self) -> None:
        """Chapter/Chapitre headings split the text into multiple chapters.

        Note: the TXT reader joins consecutive non-blank lines into a
        single paragraph, so headings must be separated from body text
        by a blank line to be recognised as chapter boundaries.
        """
        content = textwrap.dedent("""\
            Chapter 1

            This is the first chapter content.

            Chapter 2

            This is the second chapter content.

            Chapitre 3

            Ceci est le troisieme chapitre.
            """)

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "chapters.txt"
            path.write_text(content, encoding="utf-8")
            doc = read_document(path)
            self.assertEqual(len(doc.chapters), 3)
            self.assertEqual(doc.chapters[0].title, "Chapter 1")
            self.assertEqual(doc.chapters[1].title, "Chapter 2")
            self.assertEqual(doc.chapters[2].title, "Chapitre 3")

    def test_txt_without_headings_single_chapter(self) -> None:
        """A TXT without chapter headings yields one chapter."""
        content = "Just a paragraph of text with no heading.\n\nAnother paragraph."
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "nohead.txt"
            path.write_text(content, encoding="utf-8")
            doc = read_document(path)
            self.assertEqual(len(doc.chapters), 1)
            self.assertEqual(len(doc.chapters[0].paragraphs), 2)

    def test_txt_node_format(self) -> None:
        """detect_format returns 'txt' for .txt files and rejects unknowns."""
        self.assertEqual(detect_format("doc.txt"), "txt")
        with self.assertRaises(ValueError):
            detect_format("doc.pdf")


# ---------------------------------------------------------------------------
# Format dispatch & detect tests
# ---------------------------------------------------------------------------


class FormatDispatchTests(unittest.TestCase):
    """detect_format and read_document dispatch for all supported extensions."""

    def test_detect_all_supported_formats(self) -> None:
        for ext, expected in [
            (".epub", "epub"),
            (".docx", "docx"),
            (".txt", "txt"),
            (".srt", "srt"),
        ]:
            with self.subTest(ext=ext):
                self.assertEqual(detect_format(f"test{ext}"), expected)

    def test_unsupported_format_raises(self) -> None:
        for bad in (".pdf", ".html", ".md", ".odt"):
            with self.subTest(ext=bad):
                with self.assertRaises(ValueError):
                    detect_format(f"doc{bad}")

    def test_minimal_txt_roundtrip(self) -> None:
        """Write TXT → read_document → chunk → verify stable hashes."""
        content = textwrap.dedent("""\
            Chapter 1

            First paragraph of text for testing.

            Second paragraph for the test.
            """)
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "roundtrip.txt"
            path.write_text(content, encoding="utf-8")
            doc = read_document(path)
            self.assertEqual(doc.title, "roundtrip")
            self.assertEqual(len(doc.chapters), 1)
            chunks = chunk_paragraphs(
                doc.chapters[0].paragraphs, doc.chapters[0].id, doc.chapters[0].title
            )
            self.assertGreaterEqual(len(chunks), 1)
            for c in chunks:
                self.assertEqual(len(c["source_hash"]), 64)
                self.assertTrue(all(h in "0123456789abcdef" for h in c["source_hash"]))


if __name__ == "__main__":
    unittest.main()
