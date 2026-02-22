import os
import pytest
import docx
from src.formats.docx_handler import DocxHandler

def create_mock_docx(file_path):
    doc = docx.Document()
    
    # 1. Paragraph with mixed bold and italic
    p1 = doc.add_paragraph()
    p1.add_run("Title with ")
    r_bold = p1.add_run("bold")
    r_bold.bold = True
    p1.add_run(" text")
    
    # 2. Paragraph with multiple formats
    p2 = doc.add_paragraph()
    p2.add_run("A simple paragraph with ")
    r_italic = p2.add_run("italic")
    r_italic.italic = True
    p2.add_run(" and ")
    r_bold2 = p2.add_run("bold")
    r_bold2.bold = True
    p2.add_run(" formatting.")
    
    # 3. Pure text
    doc.add_paragraph("Pure text without formatting.")
    
    doc.save(file_path)

def test_docx_format_preservation(tmp_path):
    docx_path = tmp_path / "test_format.docx"
    out_path = tmp_path / "out_format.docx"
    create_mock_docx(str(docx_path))
    
    handler = DocxHandler()
    
    # 1. Read the DOCX
    segments = handler.read(str(docx_path))
    
    assert len(segments) == 3
    
    # Check extraction
    assert segments[0]['source_text'] == "Title with <r0>bold</r0> text"
    assert '<r0>' in segments[0]['metadata']['tags_map']
    assert segments[0]['metadata']['tags_map']['<r0>']['bold'] == True
    
    assert segments[1]['source_text'] == "A simple paragraph with <r0>italic</r0> and <r1>bold</r1> formatting."
    assert '<r0>' in segments[1]['metadata']['tags_map']
    assert '<r1>' in segments[1]['metadata']['tags_map']
    assert segments[1]['metadata']['tags_map']['<r0>']['italic'] == True
    assert segments[1]['metadata']['tags_map']['<r1>']['bold'] == True
    
    assert segments[2]['source_text'] == "Pure text without formatting."
    assert not segments[2]['metadata'].get('tags_map')
    
    # 2. Simulate translation
    segments[0]['target_text'] = "Titre avec texte <r0>gras</r0>"
    segments[1]['target_text'] = "Un paragraphe simple avec format <r0>italique</r0> et <r1>gras</r1>."
    segments[2]['target_text'] = "Texte pur sans formatage."
    
    # 3. Write DOCX
    handler.write(str(out_path), segments, original_file_path=str(docx_path))
    
    # 4. Read back and verify runs
    doc_out = docx.Document(str(out_path))
    
    assert len(doc_out.paragraphs) == 3
    
    # Para 1
    runs1 = doc_out.paragraphs[0].runs
    assert len(runs1) == 2
    assert runs1[0].text == "Titre avec texte "
    assert runs1[1].text == "gras"
    assert runs1[1].bold == True
    
    # Para 2
    runs2 = doc_out.paragraphs[1].runs
    assert len(runs2) == 5
    assert runs2[0].text == "Un paragraphe simple avec format "
    assert runs2[1].text == "italique"
    assert runs2[1].italic == True
    assert runs2[2].text == " et "
    assert runs2[3].text == "gras"
    assert runs2[3].bold == True
    assert runs2[3].italic in (False, None)
    assert runs2[4].text == "."
    
    # Para 3
    runs3 = doc_out.paragraphs[2].runs
    assert len(runs3) == 1
    assert runs3[0].text == "Texte pur sans formatage."
