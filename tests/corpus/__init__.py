"""Evaluation corpus for translation quality measurement.

Each extract represents a literary domain: dialogue, exposition,
description, notes, subtitles, and fantasy terms. The module provides
helpers to generate fixture files and evaluate structural integrity
(chunk → assemble → verify no text loss, order preserved).
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any


# ---------- representative extracts ----------


CORPUS_DIALOGUE = """Chapter 1

"I can't believe you did that," she said, her voice trembling.

"I had no choice." He looked away. "The gate was already opening."

"But you promised! You stood right there and said—" She stopped, biting her lip.

A long silence filled the room. Outside, the wind rattled the shutters.
"""

CORPUS_EXPOSITION = """Chapter 2

The city of Aldervale had stood for three thousand years, its walls
built from stone quarried in the Age of Kings. Scholars disagreed on
the exact founding date — some pointed to the Treaty of the Four Rivers,
others to the coronation of Alaric the First — but all agreed that no
army had ever breached the Outer Gate.

Until today.
"""

CORPUS_DESCRIPTION = """Chapter 3

The garden stretched toward the horizon, a riot of crimson roses and
silver-leafed lavender. Stone paths wound between ancient oaks whose
branches formed a canopy so dense that only slivers of afternoon light
reached the moss below. A fountain stood at the center — a marble
nymph pouring an endless stream from a tilted amphora — and around it
hummed a thousand bees drunk on nectar.
"""

CORPUS_FANTASY_TERMS = """Chapter 4

The Veil-Walker raised his staff and whispered the Binding Words.
A ripple passed through the air, and the lesser demons — the skritch,
the hollow-men, the shade-hounds — fell silent. Only the Arch-Fiend
remained, its seven eyes fixed on the sorcerer.

"You cannot hold the Veil forever, mortal," it hissed.

"Perhaps not," said Talric. "But I only need to hold it long enough."
"""

CORPUS_SUBTITLES = """Chapter 5

Run!

I said run, not negotiate.

The alarm will reset in ninety seconds.
"""

CORPUS_LONG_TEXT = ("Chapter 6\n\n" + "\n\n".join(
    f"Paragraph {i}: Mara crossed the bridge, counted the lanterns, "
    f"and repeated Talric's warning so she would not forget it."
    for i in range(1, 80)
))


# ---------- all extracts ----------


ALL_EXTRACTS: dict[str, str] = {
    "dialogue": CORPUS_DIALOGUE,
    "exposition": CORPUS_EXPOSITION,
    "description": CORPUS_DESCRIPTION,
    "fantasy_terms": CORPUS_FANTASY_TERMS,
    "subtitles": CORPUS_SUBTITLES,
    "long_text": CORPUS_LONG_TEXT,
}


@dataclass(frozen=True)
class CorpusCase:
    case_id: str
    source_format: str
    extract_key: str
    terms: tuple[str, ...] = ()
    notes: str = ""


CORPUS_CASES: tuple[CorpusCase, ...] = (
    CorpusCase("dialogue_txt", "txt", "dialogue", ("gate",), "dialogue beats"),
    CorpusCase("exposition_txt", "txt", "exposition", ("Aldervale",)),
    CorpusCase("description_docx", "docx", "description", ("crimson roses",)),
    CorpusCase(
        "fantasy_epub",
        "epub",
        "fantasy_terms",
        ("Veil-Walker", "Binding Words", "Arch-Fiend", "Talric"),
    ),
    CorpusCase("subtitles_srt", "srt", "subtitles", ("alarm",)),
    CorpusCase("long_text_txt", "txt", "long_text", ("Talric",), "chunking limits"),
)


# ---------- fixture generators ----------


def write_txt_fixture(extract_key: str, target_dir: Path) -> Path:
    """Write a TXT fixture file from an extract key, return path."""
    text = ALL_EXTRACTS.get(extract_key)
    if text is None:
        raise KeyError(f"Unknown extract: {extract_key!r}")
    path = target_dir / f"{extract_key}.txt"
    path.write_text(text, encoding="utf-8")
    return path


def write_srt_fixture(extract_key: str, target_dir: Path, lines_per_cue: int = 1) -> Path:
    """Write a minimal SRT fixture from extract text, one cue per paragraph."""
    text = ALL_EXTRACTS.get(extract_key)
    if text is None:
        raise KeyError(f"Unknown extract: {extract_key!r}")
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    cues: list[str] = []
    start_sec = 0
    for idx, para in enumerate(paragraphs, 1):
        cues.append(str(idx))
        end_sec = start_sec + max(2, len(para) // 15)
        cues.append(
            f"00:{start_sec // 60:02d}:{start_sec % 60:02d},000"
            f" --> "
            f"00:{end_sec // 60:02d}:{end_sec % 60:02d},000"
        )
        cues.append(para)
        cues.append("")
        start_sec = end_sec + 1
    path = target_dir / f"{extract_key}.srt"
    path.write_text("\n".join(cues), encoding="utf-8")
    return path


def write_docx_fixture(extract_key: str, target_dir: Path) -> Path:
    """Write a simple DOCX fixture with headings and paragraphs."""
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is in requirements
        raise RuntimeError("python-docx is required for DOCX corpus fixtures") from exc
    text = ALL_EXTRACTS.get(extract_key)
    if text is None:
        raise KeyError(f"Unknown extract: {extract_key!r}")
    doc = docx.Document()
    parts = [p.strip() for p in text.split("\n\n") if p.strip()]
    if parts:
        doc.add_heading(parts[0], level=1)
    for para in parts[1:]:
        doc.add_paragraph(para)
    path = target_dir / f"{extract_key}.docx"
    doc.save(str(path))
    return path


def write_epub_fixture(extract_key: str, target_dir: Path) -> Path:
    """Write a minimal EPUB fixture from one extract."""
    try:
        from ebooklib import epub
    except ImportError as exc:  # pragma: no cover - dependency is in requirements
        raise RuntimeError("EbookLib is required for EPUB corpus fixtures") from exc
    text = ALL_EXTRACTS.get(extract_key)
    if text is None:
        raise KeyError(f"Unknown extract: {extract_key!r}")
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    title = paragraphs[0] if paragraphs else extract_key
    body = "".join(f"<p>{p}</p>" for p in paragraphs[1:])
    book = epub.EpubBook()
    book.set_identifier(f"noveltrad-corpus-{extract_key}")
    book.set_title(title)
    book.set_language("en")
    chapter = epub.EpubHtml(title=title, file_name="chapter.xhtml", lang="en")
    chapter.content = f"<html><body><h1>{title}</h1>{body}</body></html>"
    book.add_item(chapter)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", chapter]
    book.toc = (chapter,)
    path = target_dir / f"{extract_key}.epub"
    epub.write_epub(str(path), book)
    return path


def write_case_fixture(case: CorpusCase, target_dir: Path) -> Path:
    if case.source_format == "txt":
        return write_txt_fixture(case.extract_key, target_dir)
    if case.source_format == "srt":
        return write_srt_fixture(case.extract_key, target_dir)
    if case.source_format == "docx":
        return write_docx_fixture(case.extract_key, target_dir)
    if case.source_format == "epub":
        return write_epub_fixture(case.extract_key, target_dir)
    raise ValueError(f"Unsupported corpus format: {case.source_format}")


# ---------- evaluation helpers ----------


def structural_metrics(original: str, reconstructed: str) -> dict[str, Any]:
    """Compare original and reconstructed text for structural integrity.

    Return:
        chars_lost: number of characters missing from reconstructed
        chars_added: characters present in reconstructed but not original
        word_count_delta: difference in word count
        lines_count_delta: difference in line count
    """
    orig_words = original.split()
    recon_words = reconstructed.split()
    orig_chars = len(original)
    recon_chars = len(reconstructed)
    return {
        "chars_lost": max(0, orig_chars - recon_chars),
        "chars_added": max(0, recon_chars - orig_chars),
        "word_count_original": len(orig_words),
        "word_count_reconstructed": len(recon_words),
        "word_count_delta": len(recon_words) - len(orig_words),
        "line_count_original": len(original.splitlines()),
        "line_count_reconstructed": len(reconstructed.splitlines()),
    }


def terminology_coherence(
    source_text: str, translated_text: str, terms: list[str]
) -> dict[str, Any]:
    """Check if named entities / terms are preserved in the translation."""
    found: dict[str, bool] = {}
    for term in terms:
        in_source = term in source_text
        in_target = term in translated_text
        found[term] = {"in_source": in_source, "in_target": in_target, "preserved": in_target}
    all_preserved = all(v["preserved"] for v in found.values())
    return {"terms": found, "all_preserved": all_preserved, "term_count": len(terms)}
