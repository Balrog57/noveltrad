from src.formats.format_handler import FormatHandler
import docx
import os
import re
import json

class DocxHandler(FormatHandler):
    def read(self, file_path):
        doc = docx.Document(file_path)
        segments = []
        global_index = 0
        
        for i, para in enumerate(doc.paragraphs):
            # If paragraph has no text, skip
            if not para.text.strip():
                continue
                
            source_text = ""
            run_styles = {}
            run_counter = 0
            
            for run_index, run in enumerate(para.runs):
                text = run.text
                if not text:
                    continue
                
                # Check if run has any special formatting we want to preserve
                has_formatting = any([run.bold, run.italic, run.underline, run.font.strike, run.style.name != 'Default Paragraph Font'])
                
                if has_formatting:
                    tag_name = f"r{run_counter}"
                    source_text += f"<{tag_name}>{text}</{tag_name}>"
                    
                    # Save run properties
                    run_styles[f"<{tag_name}>"] = {
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'strike': run.font.strike,
                        'style_name': run.style.name if run.style else None
                    }
                    run_counter += 1
                else:
                    source_text += text
                    
            source_text = source_text.strip()
            if source_text:
                segments.append({
                    'index': global_index,
                    'source_text': source_text,
                    'metadata': {
                        'para_index': i,
                        'tags_map': run_styles
                    }
                })
                global_index += 1
                
        return segments

    def write(self, file_path, segments, original_file_path=None):
        if not original_file_path:
             raise ValueError("Original DOCX required")

        doc = docx.Document(original_file_path)
        
        # Map segments
        segments_map = {}
        for s in segments:
            if not s.get('target_text'):
                continue
            meta = s.get('metadata')
            if not meta:
                continue
            
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except:
                    pass
            segments_map[meta['para_index']] = {
                'text': s['target_text'],
                'tags_map': meta.get('tags_map', {})
            }
        
        # Regex to find <rX> tags
        tag_pattern = re.compile(r'(<r\d+>)(.*?)(</r\d+>)', re.DOTALL)
        
        for i, para in enumerate(doc.paragraphs):
            if i in segments_map:
                new_text = segments_map[i]['text']
                tags_map = segments_map[i]['tags_map']
                
                # Clear existing paragraph runs to replace them
                para.clear()
                
                # Parse the target text to extract <rN> blocks and plain text
                # We can do this by splitting or iterating over matches
                current_pos = 0
                for match in tag_pattern.finditer(new_text):
                    # Add preceding plain text
                    plain_text = new_text[current_pos:match.start()]
                    if plain_text:
                        para.add_run(plain_text)
                        
                    open_tag = match.group(1)
                    inner_text = match.group(2)
                    
                    # Add formatted run
                    new_run = para.add_run(inner_text)
                    
                    # Apply formatting if we found it in context
                    if open_tag in tags_map:
                        style_data = tags_map[open_tag]
                        if style_data.get('bold'): new_run.bold = True
                        if style_data.get('italic'): new_run.italic = True
                        if style_data.get('underline'): new_run.underline = True
                        if style_data.get('strike'): new_run.font.strike = True
                        # Restoring style is trickier because the style object needs to exist in the doc
                        # For simple formatting, setting bold/italic directly is sufficient
                        
                    current_pos = match.end()
                    
                # Add trailing plain text
                remaining_text = new_text[current_pos:]
                if remaining_text:
                    para.add_run(remaining_text)
                
        doc.save(file_path)

    def get_supported_extensions(self):
        return ['.docx']
