"""Assembler agent — Agent 9 (last stage) of the v4 pipeline.

Reconstructs the target document from the polished chunks. Supports
EPUB, DOCX, TXT, and SRT. The output path is provided in the task
payload (under `output_path`); the orchestrator also sets it in the
project context.

The assembler only acts on the FINAL `done` for the project. It
collects all polished chunks in the project, sorts them by chapter
then chunk_index, and writes them out.
"""

from __future__ import annotations

import json
import logging
import shutil
from pathlib import Path
from typing import Any

from .base_worker import BaseWorker
from ..formats import read_project_manifest

logger = logging.getLogger(__name__)


class Worker(BaseWorker):
    use_control_thread = True

    def handle_task(self, msg: dict[str, Any]) -> dict[str, Any] | None:
        action = msg.get("action")
        if action not in ("assemble",):
            return self._emit_error(
                msg.get("chunk_id"),
                "unknown_action",
                f"assembler: unknown action {action!r}",
            )
        payload = msg.get("payload") or {}
        chunk_id = msg.get("chunk_id")
        # The assembler task carries a snapshot of the project state.
        chunks: list[dict[str, Any]] = payload.get("chunks") or []
        output_path = payload.get("output_path")
        fmt = (payload.get("format") or "txt").lower()
        if not output_path:
            return self._emit_error(
                chunk_id, "missing_output_path", "assembler: output_path required"
            )
        out = Path(str(output_path))
        out.parent.mkdir(parents=True, exist_ok=True)
        groups = _grouped_by_source(chunks)
        written: list[str] = []
        try:
            for idx, (source_file, group_chunks) in enumerate(groups):
                target = _derive_output_path(out, source_file, idx)
                if fmt == "epub" or fmt == "epub_bilingual":
                    _write_epub(
                        target,
                        group_chunks,
                        title=payload.get("title", target.stem),
                        manifest_path=payload.get("manifest_path"),
                        bilingual=(fmt == "epub_bilingual"),
                    )
                elif fmt == "docx":
                    _write_docx(target, group_chunks)
                elif fmt == "srt":
                    _write_srt(target, group_chunks)
                else:
                    _write_txt(target, group_chunks)
                written.append(str(target))
        except Exception as exc:
            logger.exception("assembler: write failed")
            return self._emit_error(chunk_id, "write_failed", str(exc))
        primary = written[0] if written else str(out)
        return self._emit_done(
            chunk_id,
            {
                "output_path": primary,
                "chunk_count": len(chunks),
                "status": "assembled",
                "outputs": written,
            },
        )


def _sorted_chunks(chunks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    def key(c: dict[str, Any]) -> tuple[str, int]:
        return (str(c.get("chapter_id") or ""), int(c.get("chunk_index") or 0))

    return sorted(chunks, key=key)


def _grouped_by_source(
    chunks: list[dict[str, Any]],
) -> list[tuple[str, list[dict[str, Any]]]]:
    """Group chunks by their ``source_file`` attribute, preserving order.

    Chunks without a ``source_file`` end up in a single bucket with an
    empty key, which is the legacy single-file behaviour. Within each
    bucket the chunks are sorted by chapter and chunk_index so the
    output is deterministic.
    """
    buckets: dict[str, list[dict[str, Any]]] = {}
    for c in chunks:
        key = str(c.get("source_file") or "")
        buckets.setdefault(key, []).append(c)
    return [(k, _sorted_chunks(v)) for k, v in buckets.items()]


def _derive_output_path(
    template: Path, source_file: str, idx: int
) -> Path:
    """If the payload holds chunks from multiple source files, append
    a per-source suffix to avoid clobbering siblings."""
    if not source_file:
        return template
    stem = Path(source_file).stem or f"file{idx}"
    suffix = template.suffix or ".txt"
    return template.with_name(f"{template.stem}_{stem}{suffix}")


def _polished_text(c: dict[str, Any]) -> str:
    return (
        c.get("polished_translation")
        or c.get("grammar_checked")
        or c.get("qa_checked")
        or c.get("glossary_applied")
        or c.get("raw_translation")
        or ""
    )


def _write_txt(path: Path, chunks: list[dict[str, Any]]) -> None:
    lines: list[str] = []
    for c in _sorted_chunks(chunks):
        title = c.get("chapter_title") or c.get("chapter_id") or ""
        if title and (not lines or lines[-1] != title):
            lines.append(f"## {title}")
            lines.append("")
        lines.append(_polished_text(c))
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")


def _write_epub(
    path: Path,
    chunks: list[dict[str, Any]],
    title: str,
    manifest_path: str | None = None,
    bilingual: bool = False,
) -> None:
    if manifest_path:
        try:
            manifest = read_project_manifest(manifest_path)
            if manifest.get("source_format") == "epub":
                _write_epub_from_manifest(path, chunks, manifest, bilingual=bilingual)
                return
        except Exception:
            logger.exception("assembler: manifest EPUB write failed; using fallback")
    try:
        from ebooklib import epub
    except ImportError as exc:
        raise RuntimeError("EPUB assembly requires EbookLib") from exc
    book = epub.EpubBook()
    book.set_identifier("noveltrad-assembled")
    book.set_title(title)
    book.set_language("en")
    chapters: list[Any] = []
    current_chap_title: str | None = None
    current_paragraphs: list[str] = []
    chapter_index = 0
    bilingual_css = ""
    if bilingual:
        bilingual_css = (
            "<style>p.original { color: #666; font-style: italic; margin-bottom: 0.2em; } "
            "p.translation { margin-top: 0; margin-bottom: 1em; }</style>"
        )

    def _flush() -> None:
        nonlocal chapter_index, current_paragraphs, current_chap_title
        if not current_paragraphs:
            return
        chapter_index += 1
        c = epub.EpubHtml(
            title=current_chap_title or f"Chapter {chapter_index}",
            file_name=f"chap_{chapter_index:04d}.xhtml",
            lang="en",
        )
        body = "".join(f"<p>{p}</p>" for p in current_paragraphs if p.strip())
        c.content = (
            f"<html><head><title>{current_chap_title or ''}</title>{bilingual_css}</head>"
            f"<body><h1>{current_chap_title or ''}</h1>{body}</body></html>"
        )
        book.add_item(c)
        chapters.append(c)
        current_paragraphs = []

    for c in _sorted_chunks(chunks):
        title_c = c.get("chapter_title") or c.get("chapter_id") or ""
        if current_chap_title is None:
            current_chap_title = title_c
        if title_c and title_c != current_chap_title:
            _flush()
            current_chap_title = title_c
        if bilingual:
            src = c.get("source_text") or ""
            tgt = _polished_text(c)
            current_paragraphs.append(
                f'<p class="original">{_escape_xml(src)}</p>'
                f'<p class="translation">{_escape_xml(tgt)}</p>'
            )
        else:
            current_paragraphs.append(_polished_text(c))
    _flush()
    book.toc = tuple(chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", *chapters]
    epub.write_epub(str(path), book)


def _write_epub_from_manifest(
    path: Path, chunks: list[dict[str, Any]], manifest: dict[str, Any],
    bilingual: bool = False,
) -> None:
    try:
        import ebooklib
        from bs4 import BeautifulSoup
        from ebooklib import epub
    except ImportError as exc:
        raise RuntimeError("EPUB assembly requires EbookLib and BeautifulSoup4") from exc

    source_epub = Path(manifest.get("working_source_path") or manifest.get("source_path"))
    if not source_epub.exists():
        raise RuntimeError(f"Manifest source EPUB not found: {source_epub}")

    book = epub.read_epub(str(source_epub), options={"ignore_ncx": True})
    translations_by_anchor: dict[tuple[str, int], list[str]] = {}
    for chunk in _sorted_chunks(chunks):
        translated = _polished_text(chunk).strip()
        if not translated:
            continue
        anchors = (chunk.get("metadata") or {}).get("anchors") or []
        if not anchors:
            continue
        first = anchors[0]
        if first.get("kind") != "epub_text_node":
            continue
        key = (str(first.get("item_id")), int(first.get("node_index", 0)))
        translations_by_anchor.setdefault(key, []).append(translated)
        for extra in anchors[1:]:
            if extra.get("kind") == "epub_text_node":
                extra_key = (
                    str(extra.get("item_id")),
                    int(extra.get("node_index", 0)),
                )
                translations_by_anchor.setdefault(extra_key, [])

    for item_id, replacements in _group_replacements(translations_by_anchor).items():
        item = book.get_item_with_id(item_id)
        if item is None or item.get_type() != ebooklib.ITEM_DOCUMENT:
            continue
        soup = BeautifulSoup(item.get_content(), "html.parser")
        if bilingual:
            _apply_bilingual_replacements(soup, replacements)
        else:
            editable_nodes = _find_editable_nodes(soup)
            for node_index, value in replacements.items():
                if 0 <= node_index < len(editable_nodes):
                    editable_nodes[node_index].replace_with(value)
        item.set_content(str(soup).encode("utf-8"))

    doc_items = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            if not getattr(item, "uid", None):
                item.uid = getattr(item, "id", None) or getattr(item, "file_name", "chapter")
            doc_items.append(item)
    if doc_items:
        book.toc = tuple(doc_items)
        book.spine = ["nav", *doc_items]
    path.parent.mkdir(parents=True, exist_ok=True)
    epub.write_epub(str(path), book)


def _find_editable_nodes(soup) -> list:
    editable_nodes = []
    for node in soup.find_all(string=True):
        normalized = " ".join(str(node).split())
        if not normalized:
            continue
        parent_name = getattr(getattr(node, "parent", None), "name", "") or ""
        if parent_name.lower() in {"script", "style", "title"}:
            continue
        editable_nodes.append(node)
    return editable_nodes


def _apply_bilingual_replacements(soup, replacements: dict[int, str]) -> None:
    from bs4 import BeautifulSoup

    editable_nodes = _find_editable_nodes(soup)
    for node_index, translation in replacements.items():
        if 0 <= node_index < len(editable_nodes):
            orig = editable_nodes[node_index]
            orig_text = str(orig)
            wrapper = soup.new_tag("p", **{"class": "original"})
            wrapper.string = orig_text
            trans_wrapper = soup.new_tag("p", **{"class": "translation"})
            trans_wrapper.string = translation
            orig.replace_with(wrapper)
            wrapper.insert_after(trans_wrapper)


def _escape_xml(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _group_replacements(
    translations_by_anchor: dict[tuple[str, int], list[str]]
) -> dict[str, dict[int, str]]:
    grouped: dict[str, dict[int, str]] = {}
    for (item_id, node_index), parts in translations_by_anchor.items():
        grouped.setdefault(item_id, {})[node_index] = "\n\n".join(parts)
    return grouped


def _write_docx(path: Path, chunks: list[dict[str, Any]]) -> None:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("DOCX assembly requires python-docx") from exc
    doc = docx.Document()
    current_chap: str | None = None
    for c in _sorted_chunks(chunks):
        title = c.get("chapter_title") or c.get("chapter_id") or ""
        if title and title != current_chap:
            doc.add_heading(title, level=1)
            current_chap = title
        text = _polished_text(c)
        if text:
            doc.add_paragraph(text)
    doc.save(str(path))


_SRT_TIME = "%H:%M:%S,%f"


def _write_srt(path: Path, chunks: list[dict[str, Any]]) -> None:
    out: list[str] = []
    index = 1
    start = 0
    for c in _sorted_chunks(chunks):
        text = _polished_text(c).strip()
        if not text:
            continue
        # Rough heuristic: 3 sec per chunk, 12 chars/sec reading speed.
        import datetime as _dt

        s = _dt.datetime.utcfromtimestamp(start)
        end_ts = start + max(3, len(text) // 12)
        e = _dt.datetime.utcfromtimestamp(end_ts)
        out.append(str(index))
        out.append(
            f"{s.strftime(_SRT_TIME)[:-3]} --> {e.strftime(_SRT_TIME)[:-3]}"
        )
        out.append(text)
        out.append("")
        index += 1
        start = end_ts + 1
    path.write_text("\n".join(out), encoding="utf-8")


__all__ = ["Worker"]
