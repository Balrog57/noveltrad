"""File format handlers — extract structured chunks from input files.

Each handler exposes a common surface:

    Handler.from_path(path) -> Document
    Document.title, .author
    Document.chapters: list[Chapter]
    Chapter.id, .title, .paragraphs: list[str]

Chunking rules (shared with the Parser agent):
  * Paragraphs < MIN_CHARS are merged with the next one to avoid
    under-sized chunks.
  * Paragraphs > MAX_CHARS are split on sentence boundaries.
  * Resulting chunks are roughly MIN_CHARS..MAX_CHARS long.

This module is GUI-free and PyQt6-free. It can be imported from the
FastAPI process and from the Parser subprocess.
"""

from __future__ import annotations

import json
import logging
import re
import shutil
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

logger = logging.getLogger(__name__)


MIN_CHARS = 200
MAX_CHARS = 800
SOFT_TARGET = 500


@dataclass
class Chapter:
    id: str
    title: str
    paragraphs: list[str] = field(default_factory=list)
    blocks: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    def plain_text(self) -> str:
        return "\n\n".join(p for p in self.paragraphs if p)


@dataclass
class Document:
    title: str
    author: str
    chapters: list[Chapter] = field(default_factory=list)
    source_format: str = "unknown"
    metadata: dict[str, Any] = field(default_factory=dict)


MANIFEST_DIR = ".noveltrad"
MANIFEST_NAME = "project_manifest.json"


# ---------- chunking ----------


def chunk_paragraphs(
    paragraphs: Iterable[str],
    chapter_id: str,
    chapter_title: str,
    min_chars: int = MIN_CHARS,
    max_chars: int = MAX_CHARS,
    soft_target: int = SOFT_TARGET,
) -> list[dict[str, Any]]:
    """Group paragraphs into chunks of roughly soft_target characters.

    The output list contains dicts with: id, chapter_id, chapter_title,
    chunk_index, source_text, source_hash. `id` is a deterministic
    sha1-based hash so re-parsing yields stable IDs (resumability).
    """
    blocks = [{"text": p, "anchors": []} for p in paragraphs]
    return chunk_blocks(
        blocks,
        chapter_id=chapter_id,
        chapter_title=chapter_title,
        min_chars=min_chars,
        max_chars=max_chars,
        soft_target=soft_target,
    )


def chunk_blocks(
    blocks: Iterable[dict[str, Any]],
    chapter_id: str,
    chapter_title: str,
    min_chars: int = MIN_CHARS,
    max_chars: int = MAX_CHARS,
    soft_target: int = SOFT_TARGET,
) -> list[dict[str, Any]]:
    """Group text blocks while preserving format anchors in chunk metadata."""
    import hashlib

    merged: list[dict[str, Any]] = []
    buf_text = ""
    buf_anchors: list[dict[str, Any]] = []
    source_block_indexes: list[int] = []

    def _flush() -> None:
        nonlocal buf_text, buf_anchors, source_block_indexes
        text = buf_text.strip()
        if text:
            merged.append(
                {
                    "text": text,
                    "anchors": list(buf_anchors),
                    "source_block_indexes": list(source_block_indexes),
                }
            )
        buf_text = ""
        buf_anchors = []
        source_block_indexes = []

    for block_index, block in enumerate(blocks):
        text = (block.get("text") or "").strip()
        if not text:
            continue
        anchors = list(block.get("anchors") or [])
        if not buf_text:
            buf_text = text
        elif len(buf_text) < min_chars or len(buf_text) + len(text) <= soft_target:
            buf_text = f"{buf_text}\n\n{text}"
        else:
            _flush()
            buf_text = text
        buf_anchors.extend(anchors)
        source_block_indexes.append(block_index)
    _flush()

    out: list[dict[str, Any]] = []
    chunk_index = 0
    for merged_block in merged:
        text = merged_block["text"]
        parts = _split_oversize(text, max_chars)
        for part_index, sub in enumerate(parts):
            chunk_id = _stable_chunk_id(chapter_id, chunk_index, sub)
            out.append(
                {
                    "id": chunk_id,
                    "chapter_id": chapter_id,
                    "chapter_title": chapter_title,
                    "chunk_index": chunk_index,
                    "source_text": sub,
                    "source_hash": hashlib.sha256(sub.encode("utf-8")).hexdigest(),
                    "metadata": {
                        "anchors": merged_block.get("anchors") or [],
                        "source_block_indexes": merged_block.get("source_block_indexes") or [],
                        "split_part": part_index,
                        "split_total": len(parts),
                    },
                }
            )
            chunk_index += 1
    return out


def _split_oversize(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    parts: list[str] = []
    cursor = 0
    while cursor < len(text):
        if cursor + max_chars >= len(text):
            parts.append(text[cursor:].strip())
            break
        window = text[cursor : cursor + max_chars]
        cut = _find_sentence_break(window)
        if cut is None or cut < max_chars // 2:
            cut = max_chars
        parts.append(text[cursor : cursor + cut].strip())
        cursor += cut
    return [p for p in parts if p]


_SENTENCE_END = re.compile(r"[.!?。!?][\"')\]]?\s")


def _find_sentence_break(window: str) -> int | None:
    last = None
    for m in _SENTENCE_END.finditer(window):
        last = m.end()
    return last


def _stable_chunk_id(chapter_id: str, idx: int, text: str) -> str:
    import hashlib

    h = hashlib.sha1(f"{chapter_id}::{idx}::{text}".encode("utf-8")).hexdigest()
    return h[:16]


# ---------- format-specific extractors ----------


def _read_epub(path: Path) -> Document:
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise RuntimeError(
            "EPUB support requires EbookLib and BeautifulSoup4"
        ) from exc

    book = epub.read_epub(str(path), options={"ignore_ncx": True})
    title = _safe_meta(book, "title") or path.stem
    author = _safe_meta(book, "creator") or _safe_meta(book, "author") or ""
    chapters: list[Chapter] = []
    spine_ids = list(book.spine)
    spine_payload: list[dict[str, Any]] = []
    for i, spine_ref in enumerate(spine_ids):
        item_id, _ = spine_ref
        item = book.get_item_with_id(item_id)
        if item is None or item.get_type() != ebooklib.ITEM_DOCUMENT:
            continue
        soup = BeautifulSoup(item.get_content(), "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        blocks: list[dict[str, Any]] = []
        node_index = 0
        for node in soup.find_all(string=True):
            if _is_markup_control_node(node):
                continue
            text = str(node)
            normalized = " ".join(text.split())
            if not normalized:
                continue
            parent_name = getattr(getattr(node, "parent", None), "name", "") or ""
            if parent_name.lower() in {"script", "style", "title"}:
                continue
            anchor = {
                "kind": "epub_text_node",
                "item_id": item_id,
                "href": getattr(item, "file_name", ""),
                "node_index": node_index,
            }
            blocks.append({"text": normalized, "anchors": [anchor]})
            node_index += 1
        if not blocks:
            continue
        paragraphs = [b["text"] for b in blocks]
        chap_title = _first_meaningful_line(paragraphs) or f"Chapter {i + 1}"
        spine_payload.append(
            {
                "id": item_id,
                "href": getattr(item, "file_name", ""),
                "chapter_id": f"ch{i + 1:04d}",
                "title": chap_title,
                "text_nodes": node_index,
            }
        )
        chapters.append(
            Chapter(
                id=f"ch{i + 1:04d}",
                title=chap_title,
                paragraphs=paragraphs,
                blocks=blocks,
                metadata={"epub_item_id": item_id},
            )
        )
    return Document(
        title=title,
        author=author,
        chapters=chapters,
        source_format="epub",
        metadata={"path": str(path), "epub_spine": spine_payload},
    )


def _safe_meta(book: Any, key: str) -> str | None:
    try:
        data = book.get_metadata("DC", key)
        if not data:
            return None
        return str(data[0][0]) if data else None
    except Exception:
        return None


def _first_meaningful_line(paragraphs: list[str]) -> str | None:
    for p in paragraphs[:3]:
        if 3 <= len(p) <= 120 and not p.endswith("."):
            return p
    return None


def _is_markup_control_node(node: Any) -> bool:
    try:
        from bs4.element import Declaration, Doctype, ProcessingInstruction
    except ImportError:
        return False
    return isinstance(node, (Declaration, Doctype, ProcessingInstruction))


def _read_docx(path: Path) -> Document:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError("DOCX support requires python-docx") from exc

    doc = docx.Document(str(path))
    title = doc.core_properties.title or path.stem
    author = doc.core_properties.author or ""
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    chapters = [
        Chapter(
            id="ch0001",
            title=title or "Document",
            paragraphs=paragraphs,
        )
    ]
    return Document(
        title=title,
        author=author,
        chapters=chapters,
        source_format="docx",
        metadata={"path": str(path)},
    )


def _read_txt(path: Path) -> Document:
    text = path.read_text(encoding="utf-8", errors="replace")
    lines = [ln.rstrip() for ln in text.splitlines()]
    paragraphs: list[str] = []
    buf: list[str] = []
    for ln in lines:
        if not ln.strip():
            if buf:
                paragraphs.append(" ".join(buf).strip())
                buf = []
            continue
        buf.append(ln.strip())
    if buf:
        paragraphs.append(" ".join(buf).strip())

    chapters: list[Chapter] = []
    current: Chapter | None = None
    heading_re = re.compile(r"^(chapter|chapitre|第[0-9一二三四五六七八九十百]+章)\b", re.I)
    for idx, p in enumerate(paragraphs):
        if heading_re.match(p) and len(p) < 120:
            if current is not None:
                chapters.append(current)
            current = Chapter(
                id=f"ch{len(chapters) + 1:04d}",
                title=p,
                paragraphs=[],
            )
        else:
            if current is None:
                current = Chapter(
                    id="ch0001",
                    title=path.stem,
                    paragraphs=[],
                )
            current.paragraphs.append(p)
    if current is not None:
        chapters.append(current)
    if not chapters:
        chapters.append(Chapter(id="ch0001", title=path.stem, paragraphs=paragraphs))
    return Document(
        title=path.stem,
        author="",
        chapters=chapters,
        source_format="txt",
        metadata={"path": str(path)},
    )


_SRT_TIME = re.compile(
    r"(\d{2}):(\d{2}):(\d{2})[,\.](\d{3})\s*-->\s*(\d{2}):(\d{2}):(\d{2})[,\.](\d{3})"
)


def _read_srt(path: Path) -> Document:
    text = path.read_text(encoding="utf-8", errors="replace")
    blocks = re.split(r"\n\s*\n", text.strip())
    paragraphs: list[str] = []
    for block in blocks:
        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
        if not lines:
            continue
        if lines[0].isdigit():
            lines = lines[1:]
        if not lines:
            continue
        m = _SRT_TIME.match(lines[0]) if lines else None
        if m:
            lines = lines[1:]
        paragraphs.append(" ".join(lines))
    return Document(
        title=path.stem,
        author="",
        chapters=[Chapter(id="ch0001", title=path.stem, paragraphs=paragraphs)],
        source_format="srt",
        metadata={"path": str(path)},
    )


_READERS = {
    ".epub": _read_epub,
    ".docx": _read_docx,
    ".txt": _read_txt,
    ".srt": _read_srt,
}


def detect_format(path: str | Path) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in _READERS:
        return suffix.lstrip(".")
    raise ValueError(f"Unsupported file format: {suffix!r}")


def read_document(path: str | Path) -> Document:
    p = Path(path)
    suffix = p.suffix.lower()
    reader = _READERS.get(suffix)
    if reader is None:
        raise ValueError(f"Unsupported file format: {suffix!r}")
    return reader(p)


def prepare_project_source(
    source_path: str | Path,
    project_dir: str | Path,
    project_id: str,
    output_path: str | Path | None = None,
) -> dict[str, Any]:
    """Copy the source into a project folder and create a base manifest."""
    source = Path(source_path).resolve()
    project = Path(project_dir).resolve()
    source_dir = project / "source"
    target_dir = project / "target"
    state_dir = project / MANIFEST_DIR
    source_dir.mkdir(parents=True, exist_ok=True)
    target_dir.mkdir(parents=True, exist_ok=True)
    state_dir.mkdir(parents=True, exist_ok=True)
    working = source_dir / source.name
    if source != working:
        shutil.copy2(source, working)
    fmt = detect_format(working)
    target = Path(output_path) if output_path else target_dir / f"{source.stem}.{fmt}"
    manifest = {
        "project_id": project_id,
        "source_format": fmt,
        "source_path": str(source),
        "working_source_path": str(working),
        "target_path": str(target),
        "chapters": [],
        "chunks": [],
        "format_payload": {},
    }
    return manifest


def write_project_manifest(project_dir: str | Path, manifest: dict[str, Any]) -> Path:
    path = Path(project_dir) / MANIFEST_DIR / MANIFEST_NAME
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def read_project_manifest(path: str | Path) -> dict[str, Any]:
    return json.loads(Path(path).read_text(encoding="utf-8"))


__all__ = [
    "Chapter",
    "Document",
    "chunk_blocks",
    "chunk_paragraphs",
    "detect_format",
    "prepare_project_source",
    "read_document",
    "read_project_manifest",
    "write_project_manifest",
    "MIN_CHARS",
    "MAX_CHARS",
    "SOFT_TARGET",
]
